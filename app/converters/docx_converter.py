"""
DOCX Converter
Converts DOCX files to PDF, TXT, HTML, MD formats.
Preserves formatting and layout structure (fonts, bold, italic, underline, list, table styles) when converting.
"""
import os
import asyncio
import subprocess
import shutil
import re


async def convert_docx(input_path: str, output_dir: str, target_format: str) -> dict:
    """DOCX converter supporting PDF, TXT, HTML, MD outputs."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _process_docx, input_path, output_dir, target_format)


def _process_docx(input_path: str, output_dir: str, target_format: str) -> dict:
    try:
        filename = os.path.basename(input_path)
        name, ext = os.path.splitext(filename)
        output_filename = f"{name}.{target_format}"
        output_path = os.path.join(output_dir, output_filename)

        if target_format == 'pdf':
            return _docx_to_pdf(input_path, output_path, output_filename)
        elif target_format == 'txt':
            return _docx_to_txt(input_path, output_path, output_filename)
        elif target_format == 'html':
            return _docx_to_html(input_path, output_path, output_filename, name)
        elif target_format == 'md':
            return _docx_to_md(input_path, output_path, output_filename, name)
        else:
            return {"success": False, "error": f"Unsupported target format for DOCX: {target_format}"}

    except Exception as e:
        return {"success": False, "error": f"DOCX conversion error: {str(e)}"}


def _docx_to_pdf(input_path: str, output_path: str, output_filename: str) -> dict:
    """Convert DOCX to PDF using docx2pdf (requires Word) or LibreOffice."""
    
    # Method 1: Try docx2pdf (Windows + Microsoft Word)
    try:
        from docx2pdf import convert
        convert(input_path, output_path)
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "filename": output_filename}
    except Exception as e:
        print(f"[DOCX->PDF] docx2pdf failed: {e}")
    
    # Method 2: Try LibreOffice
    try:
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice",  # If in PATH
            "/usr/bin/soffice",  # Linux
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"  # macOS
        ]
        
        soffice_path = None
        for path in libreoffice_paths:
            if os.path.exists(path) or shutil.which(path):
                soffice_path = path
                break
        
        if soffice_path:
            output_dir = os.path.dirname(output_path)
            result = subprocess.run([
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                input_path
            ], capture_output=True, text=True, timeout=120)
            
            expected_output = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
            if os.path.exists(expected_output):
                if expected_output != output_path:
                    shutil.move(expected_output, output_path)
                return {"success": True, "output_path": output_path, "filename": output_filename}
    except Exception as e:
        print(f"[DOCX->PDF] LibreOffice failed: {e}")
    
    # Method 3: Python-only fallback (ReportLab) - used when Word/LibreOffice are missing
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        
        doc = Document(input_path)
        pdf = SimpleDocTemplate(output_path, pagesize=letter,
                                rightMargin=40, leftMargin=40,
                                topMargin=40, bottomMargin=40)
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=10,
            leading=14
        )
        
        story = []
        
        # Process element by element in order
        body_element = doc._element.body
        p_idx = 0
        t_idx = 0
        
        # Helper to convert alignment
        def get_align_val(alignment):
            if alignment == 1:
                return 1 # Center
            elif alignment == 2:
                return 2 # Right
            elif alignment == 3:
                return 4 # Justified
            return 0 # Left
            
        def process_paragraph_runs(p):
            p_html = ""
            for run in p.runs:
                text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                if not text:
                    continue
                
                tags_start = ""
                tags_end = ""
                
                font_attrs = []
                if run.font.size:
                    font_attrs.append(f'size="{run.font.size.pt}"')
                if run.font.color and run.font.color.rgb:
                    hex_color = f"#{run.font.color.rgb}"
                    font_attrs.append(f'color="{hex_color}"')
                if run.font.name:
                    font_name = run.font.name.lower()
                    if "times" in font_name:
                        font_attrs.append('name="Times-Roman"')
                    elif "courier" in font_name:
                        font_attrs.append('name="Courier"')
                    else:
                        font_attrs.append('name="Helvetica"')
                
                if font_attrs:
                    tags_start += f'<font {" ".join(font_attrs)}>'
                    tags_end = '</font>' + tags_end
                    
                if run.bold:
                    tags_start += '<b>'
                    tags_end = '</b>' + tags_end
                if run.italic:
                    tags_start += '<i>'
                    tags_end = '</i>' + tags_end
                if run.underline:
                    tags_start += '<u>'
                    tags_end = '</u>' + tags_end
                    
                p_html += f"{tags_start}{text}{tags_end}"
            return p_html

        for child in body_element.iterchildren():
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                from docx.text.paragraph import Paragraph as DocxParagraph
                p = DocxParagraph(child, doc)
                p_html = process_paragraph_runs(p)
                
                if p_html.strip():
                    p_style = ParagraphStyle(
                        f'PStyle_{p_idx}',
                        parent=body_style,
                        alignment=get_align_val(p.alignment)
                    )
                    story.append(Paragraph(p_html, p_style))
                    p_idx += 1
                elif not p.text.strip():
                    story.append(Spacer(1, 10))
                    
            elif tag == 'tbl':
                from docx.table import Table as DocxTable
                t = DocxTable(child, doc)
                
                table_data = []
                t_styles = [
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ('LEFTPADDING', (0,0), (-1,-1), 6),
                    ('RIGHTPADDING', (0,0), (-1,-1), 6),
                ]
                
                for r_idx, row in enumerate(t.rows):
                    row_data = []
                    for c_idx, cell in enumerate(row.cells):
                        cell_paragraphs = []
                        # Cell background shading
                        try:
                            tcPr = cell._element.tcPr
                            if tcPr is not None:
                                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                if shd is not None:
                                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                                    if fill and fill != 'auto':
                                        t_styles.append(('BACKGROUND', (c_idx, r_idx), (c_idx, r_idx), colors.HexColor(f"#{fill}")))
                        except:
                            pass
                            
                        for p_idx_c, p in enumerate(cell.paragraphs):
                            p_html = process_paragraph_runs(p)
                            cp_style = ParagraphStyle(
                                f'TStyle_{t_idx}_{r_idx}_{c_idx}_{p_idx_c}',
                                parent=body_style,
                                alignment=get_align_val(p.alignment)
                            )
                            cell_paragraphs.append(Paragraph(p_html or "&nbsp;", cp_style))
                        row_data.append(cell_paragraphs)
                    table_data.append(row_data)
                
                if table_data:
                    rl_table = Table(table_data)
                    rl_table.setStyle(TableStyle(t_styles))
                    story.append(Spacer(1, 10))
                    story.append(rl_table)
                    story.append(Spacer(1, 10))
                    t_idx += 1
                
        pdf.build(story)
        if os.path.exists(output_path):
            return {
                "success": True, 
                "output_path": output_path, 
                "filename": output_filename,
                "note": "Converted using ReportLab fallback (Word/LibreOffice missing)"
            }
    except Exception as fallback_err:
        print(f"[DOCX->PDF] Fallback failed: {fallback_err}")
    
    return {
        "success": False, 
        "error": "PDF conversion failed. Microsoft Word or LibreOffice is missing and python fallback encountered an error."
    }


def _docx_to_txt(input_path: str, output_path: str, output_filename: str) -> dict:
    """Extract text from DOCX."""
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run 'pip install python-docx'"}
    
    doc = Document(input_path)
    text_content = []
    
    # Extract element by element in order
    body_element = doc._element.body
    for child in body_element.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            text_content.append(p.text)
        elif tag == 'tbl':
            from docx.table import Table
            t = Table(child, doc)
            for row in t.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join(p.text for p in cell.paragraphs).strip()
                    row_text.append(cell_text)
                text_content.append("\t".join(row_text))
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(text_content))
    
    return {"success": True, "output_path": output_path, "filename": output_filename}


def _docx_to_html(input_path: str, output_path: str, output_filename: str, name: str) -> dict:
    """Convert DOCX to HTML with high-fidelity styling preservation."""
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run 'pip install python-docx'"}
    
    doc = Document(input_path)
    
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{name}</title>
    <style>
        body {{ 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
            max-width: 800px; 
            margin: 40px auto; 
            padding: 20px; 
            line-height: 1.6; 
        }}
        h1 {{ font-size: 2em; margin-top: 1em; }}
        h2 {{ font-size: 1.5em; margin-top: 0.8em; }}
        h3 {{ font-size: 1.2em; margin-top: 0.6em; }}
        p {{ margin: 1em 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; vertical-align: top; }}
        th {{ background-color: #f5f5f5; }}
    </style>
</head>
<body>
"""
    
    def get_align_str(alignment):
        if alignment == 1:
            return "center"
        elif alignment == 2:
            return "right"
        elif alignment == 3:
            return "justify"
        return "left"

    def serialize_para(p):
        align = get_align_str(p.alignment)
        para_styles = []
        if align != "left":
            para_styles.append(f"text-align: {align}")
            
        style_str = f' style="{" ;".join(para_styles)}"' if para_styles else ""
        
        tag = "p"
        if p.style and p.style.name:
            style_name = p.style.name
            if "Heading 1" in style_name:
                tag = "h1"
            elif "Heading 2" in style_name:
                tag = "h2"
            elif "Heading 3" in style_name:
                tag = "h3"
                
        runs_html = ""
        for run in p.runs:
            text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if not text:
                continue
            
            run_styles = []
            if run.bold:
                run_styles.append("font-weight: bold")
            if run.italic:
                run_styles.append("font-style: italic")
            if run.underline:
                run_styles.append("text-decoration: underline")
            if run.font.size:
                run_styles.append(f"font-size: {run.font.size.pt}pt")
            if run.font.color and run.font.color.rgb:
                run_styles.append(f"color: #{run.font.color.rgb}")
            if run.font.name:
                run_styles.append(f"font-family: '{run.font.name}'")
                
            style_attr = f' style="{"; ".join(run_styles)}"' if run_styles else ""
            if style_attr:
                runs_html += f'<span{style_attr}>{text}</span>'
            else:
                runs_html += text
                
        if runs_html.strip():
            return f"<{tag}{style_str}>{runs_html}</{tag}>\n"
        elif not p.text.strip():
            return "<p>&nbsp;</p>\n"
        return ""

    body_element = doc._element.body
    for child in body_element.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            html_content += serialize_para(p)
        elif tag == 'tbl':
            from docx.table import Table
            t = Table(child, doc)
            html_content += "<table>\n"
            for row in t.rows:
                html_content += "  <tr>\n"
                for cell in row.cells:
                    cell_bg = ""
                    try:
                        tcPr = cell._element.tcPr
                        if tcPr is not None:
                            shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                            if shd is not None:
                                fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                                if fill and fill != 'auto':
                                    cell_bg = f' style="background-color: #{fill}"'
                    except:
                        pass
                        
                    html_content += f"    <td{cell_bg}>\n"
                    for p in cell.paragraphs:
                        html_content += serialize_para(p)
                    html_content += "    </td>\n"
                html_content += "  </tr>\n"
            html_content += "</table>\n"
            
    html_content += "</body>\n</html>"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return {"success": True, "output_path": output_path, "filename": output_filename}


def _docx_to_md(input_path: str, output_path: str, output_filename: str, name: str) -> dict:
    """Convert DOCX to Markdown with formatting preservation."""
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run 'pip install python-docx'"}
    
    doc = Document(input_path)
    md_content = f"# {name}\n\n"
    
    def serialize_para(p):
        tag = ""
        if p.style and p.style.name:
            style_name = p.style.name
            if "Heading 1" in style_name:
                tag = "# "
            elif "Heading 2" in style_name:
                tag = "## "
            elif "Heading 3" in style_name:
                tag = "### "
                
        runs_md = ""
        for run in p.runs:
            text = run.text
            if not text:
                continue
            
            if run.bold and run.italic:
                runs_md += f"***{text}***"
            elif run.bold:
                runs_md += f"**{text}**"
            elif run.italic:
                runs_md += f"*{text}*"
            elif run.underline:
                runs_md += f"<u>{text}</u>"
            else:
                runs_md += text
                
        if runs_md.strip():
            return f"{tag}{runs_md}\n\n"
        elif not p.text.strip():
            return "\n"
        return ""

    body_element = doc._element.body
    for child in body_element.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            md_content += serialize_para(p)
        elif tag == 'tbl':
            from docx.table import Table
            t = Table(child, doc)
            if t.rows:
                header_cells = t.rows[0].cells
                header_texts = []
                for cell in header_cells:
                    cell_text = " ".join(p.text for p in cell.paragraphs).strip()
                    header_texts.append(cell_text or " ")
                
                md_content += "| " + " | ".join(header_texts) + " |\n"
                md_content += "| " + " | ".join(["---"] * len(header_texts)) + " |\n"
                
                for row in t.rows[1:]:
                    cells = row.cells
                    row_texts = []
                    for cell in cells:
                        cell_text = " ".join(p.text for p in cell.paragraphs).strip()
                        row_texts.append(cell_text or " ")
                    md_content += "| " + " | ".join(row_texts) + " |\n"
                md_content += "\n"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    return {"success": True, "output_path": output_path, "filename": output_filename}
