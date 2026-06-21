"""
PDF Converter - Enhanced Version
Converts PDF files to DOCX (with formatting), TXT, HTML, MD, RTF, and images.
Uses pdf2docx for high-quality DOCX conversion, and custom PyMuPDF extraction for fallbacks.
Preserves layouts, fonts, bold, italic, and underline styling.
"""
import os
import asyncio


async def convert_pdf(input_path: str, output_dir: str, target_format: str) -> dict:
    """PDF converter supporting multiple output formats with formatting preservation."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _process_pdf, input_path, output_dir, target_format)


def _process_pdf(input_path: str, output_dir: str, target_format: str) -> dict:
    try:
        filename = os.path.basename(input_path)
        name, ext = os.path.splitext(filename)
        output_filename = f"{name}.{target_format}"
        output_path = os.path.join(output_dir, output_filename)

        if target_format in ['docx', 'doc']:
            return _pdf_to_docx(input_path, output_path, output_filename)
        elif target_format == 'txt':
            return _pdf_to_txt(input_path, output_path, output_filename)
        elif target_format == 'html':
            return _pdf_to_html(input_path, output_path, output_filename, name)
        elif target_format == 'md':
            return _pdf_to_md(input_path, output_path, output_filename, name)
        elif target_format == 'rtf':
            return _pdf_to_rtf(input_path, output_path, output_filename)
        elif target_format in ['png', 'jpg', 'jpeg']:
            return _pdf_to_images(input_path, output_dir, target_format, name)
        else:
            return {"success": False, "error": f"Unsupported target format for PDF: {target_format}"}

    except Exception as e:
        return {"success": False, "error": f"PDF conversion error: {str(e)}"}


def _pdf_to_docx(input_path: str, output_path: str, output_filename: str) -> dict:
    """Convert PDF to DOCX with formatting preservation using pdf2docx."""
    try:
        from pdf2docx import Converter
        
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "filename": output_filename}
        return {"success": False, "error": "DOCX file was not created"}
        
    except (ImportError, Exception) as e:
        # Fallback to basic text extraction if pdf2docx is not installed or fails
        return _pdf_to_docx_fallback(input_path, output_path, output_filename, str(e))


def _extract_pdf_text_with_ocr_fallback(input_path: str) -> str:
    """Helper to extract text from PDF, falling back to OCR if empty."""
    import PyPDF2
    text_content = ""
    try:
        with open(input_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
    except Exception as e:
        print(f"PyPDF2 text extraction error: {e}")
        
    if not text_content.strip():
        # Scanned PDF OCR fallback
        try:
            import fitz
            import pytesseract
            from PIL import Image
            from io import BytesIO
            
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            pdf_doc = fitz.open(input_path)
            ocr_text = []
            for page in pdf_doc:
                pix = page.get_pixmap(dpi=150)
                img_data = pix.tobytes("png")
                img = Image.open(BytesIO(img_data))
                
                page_text = ""
                try:
                    page_text = pytesseract.image_to_string(img)
                except Exception as tess_err:
                    print(f"Tesseract OCR failed on page: {tess_err}. Trying Groq Vision.")
                    temp_img_path = f"{input_path}_page_ocr_temp.png"
                    img.save(temp_img_path)
                    try:
                        from .ocr import _ocr_with_groq_vision
                        page_text = _ocr_with_groq_vision(temp_img_path, extract_style=False)
                    except Exception as groq_err:
                        print(f"Groq Vision page OCR failed: {groq_err}")
                    finally:
                        if os.path.exists(temp_img_path):
                            os.remove(temp_img_path)
                            
                if page_text:
                    ocr_text.append(page_text)
            pdf_doc.close()
            text_content = "\n\n".join(ocr_text)
        except Exception as ocr_err:
            print(f"OCR fallback error: {ocr_err}")
            
    return text_content


def _pdf_to_docx_fallback(input_path: str, output_path: str, output_filename: str, original_error: str) -> dict:
    """Fallback: High-fidelity PDF to DOCX conversion using PyMuPDF (fitz) + python-docx to preserve styling."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        import fitz  # PyMuPDF is installed on Vercel
        
        doc = Document()
        pdf_doc = fitz.open(input_path)
        
        for page in pdf_doc:
            text_blocks = page.get_text("dict")["blocks"]
            for block in text_blocks:
                if block.get("type") == 0:  # Text block
                    p = doc.add_paragraph()
                    
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if not text:
                                continue
                                
                            run = p.add_run(text)
                            
                            font_name = span.get("font", "").lower()
                            flags = span.get("flags", 0)
                            is_bold = bool(flags & 16) or "bold" in font_name or "bd" in font_name
                            is_italic = bool(flags & 2) or "italic" in font_name or "it" in font_name or "oblique" in font_name
                            
                            if is_bold:
                                run.bold = True
                            if is_italic:
                                run.italic = True
                                
                            size = span.get("size", 10)
                            run.font.size = Pt(size)
                            
                            color_int = span.get("color", 0)
                            r = (color_int >> 16) & 255
                            g = (color_int >> 8) & 255
                            b = color_int & 255
                            run.font.color.rgb = RGBColor(r, g, b)
                            
                            if "times" in font_name:
                                run.font.name = "Times New Roman"
                            elif "arial" in font_name:
                                run.font.name = "Arial"
                            elif "courier" in font_name:
                                run.font.name = "Courier New"
                            elif "calibri" in font_name:
                                run.font.name = "Calibri"
                            else:
                                run.font.name = "Arial"
                                
                        p.add_run(" ")
        
        pdf_doc.close()
        doc.save(output_path)
        
        return {
            "success": True, 
            "output_path": output_path, 
            "filename": output_filename,
            "note": "Converted using PyMuPDF high-fidelity layout fallback"
        }
    except Exception as e:
        print(f"[PDF→DOCX Fallback] High-fidelity method failed: {e}")
        return _pdf_to_docx_basic_fallback(input_path, output_path, output_filename, original_error)


def _pdf_to_docx_basic_fallback(input_path: str, output_path: str, output_filename: str, original_error: str) -> dict:
    """Basic Fallback: PDF to DOCX conversion using PyPDF2 + python-docx."""
    try:
        from docx import Document
        
        text_content = _extract_pdf_text_with_ocr_fallback(input_path)
        if not text_content.strip():
            return {"success": False, "error": f"Could not extract text from PDF. Original error: {original_error}"}
        
        doc = Document()
        for para in text_content.split('\n'):
            if para.strip():
                doc.add_paragraph(para)
        doc.save(output_path)
        
        return {
            "success": True, 
            "output_path": output_path, 
            "filename": output_filename,
            "note": "Text extraction (with OCR fallback) used for DOCX"
        }
    except Exception as e:
        return {"success": False, "error": f"PDF conversion failed: {str(e)}"}


def _pdf_to_txt(input_path: str, output_path: str, output_filename: str) -> dict:
    """Extract text from PDF with OCR fallback."""
    text_content = _extract_pdf_text_with_ocr_fallback(input_path)
    if not text_content.strip():
        return {"success": False, "error": "Could not extract text from PDF (may be image-based and OCR found no text)"}

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text_content)
    
    return {"success": True, "output_path": output_path, "filename": output_filename}


def _pdf_to_html(input_path: str, output_path: str, output_filename: str, name: str) -> dict:
    """Convert PDF to HTML preserving formatting and layout using PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(input_path)
        html_pages = []
        for page in doc:
            html_pages.append(page.get_text("html"))
        doc.close()
        
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{name}</title>
    <style>
        body {{ font-family: sans-serif; margin: 20px; }}
        .pdf-page {{ margin-bottom: 30px; border-bottom: 2px dashed #ccc; padding-bottom: 20px; }}
    </style>
</head>
<body>
"""
        for page_html in html_pages:
            full_html += f'<div class="pdf-page">\n{page_html}\n</div>\n'
        full_html += "</body>\n</html>"
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
            
        return {"success": True, "output_path": output_path, "filename": output_filename}
        
    except Exception as e:
        print(f"[PDF→HTML] PyMuPDF failed: {e}. Falling back to plain text HTML.")
        
        text_content = _extract_pdf_text_with_ocr_fallback(input_path)
        if not text_content.strip():
            return {"success": False, "error": "Could not extract text from PDF."}

        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{name}</title>
    <style>
        body {{ 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
            max-width: 800px; 
            margin: 40px auto; 
            padding: 20px; 
            line-height: 1.6; 
        }}
        p {{ margin: 1em 0; }}
    </style>
</head>
<body>
    <h1>{name}</h1>
"""
        for para in text_content.split('\n'):
            if para.strip():
                html_content += f"    <p>{para}</p>\n"
        html_content += "</body>\n</html>"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return {"success": True, "output_path": output_path, "filename": output_filename}


def _pdf_to_md(input_path: str, output_path: str, output_filename: str, name: str) -> dict:
    """Convert PDF to Markdown with high-fidelity formatting preservation."""
    try:
        import fitz
        doc = fitz.open(input_path)
        md_content = f"# {name}\n\n"
        
        for page in doc:
            text_blocks = page.get_text("dict")["blocks"]
            for block in text_blocks:
                if block.get("type") == 0:  # Text block
                    block_text = []
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if not text.strip():
                                continue
                            
                            flags = span.get("flags", 0)
                            is_bold = bool(flags & 16) or "bold" in span.get("font", "").lower()
                            is_italic = bool(flags & 2) or "italic" in span.get("font", "").lower()
                            
                            if is_bold and is_italic:
                                line_text += f" ***{text}***"
                            elif is_bold:
                                line_text += f" **{text}**"
                            elif is_italic:
                                line_text += f" *{text}*"
                            else:
                                line_text += f" {text}"
                        if line_text:
                            block_text.append(line_text.strip())
                    if block_text:
                        md_content += "\n".join(block_text) + "\n\n"
        doc.close()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
            
        return {"success": True, "output_path": output_path, "filename": output_filename}
        
    except Exception as e:
        print(f"[PDF→MD] PyMuPDF failed: {e}. Falling back to plain text Markdown.")
        
        text_content = _extract_pdf_text_with_ocr_fallback(input_path)
        if not text_content.strip():
            return {"success": False, "error": "Could not extract text from PDF."}

        md_content = f"# {name}\n\n"
        for para in text_content.split('\n'):
            if para.strip():
                md_content += f"{para}\n\n"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return {"success": True, "output_path": output_path, "filename": output_filename}


def _pdf_to_rtf(input_path: str, output_path: str, output_filename: str) -> dict:
    """Convert PDF to RTF with formatting preservation."""
    try:
        import fitz
        doc = fitz.open(input_path)
        rtf_content = "{\\rtf1\\ansi\\deff0\n"
        
        for page in doc:
            text_blocks = page.get_text("dict")["blocks"]
            for block in text_blocks:
                if block.get("type") == 0:  # Text block
                    block_text = ""
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if not text:
                                continue
                            
                            flags = span.get("flags", 0)
                            is_bold = bool(flags & 16) or "bold" in span.get("font", "").lower()
                            is_italic = bool(flags & 2) or "italic" in span.get("font", "").lower()
                            
                            # Unicode RTF escaping
                            escaped = ""
                            for char in text:
                                cp = ord(char)
                                if cp < 128:
                                    if char in ['\\', '{', '}']:
                                        escaped += '\\' + char
                                    else:
                                        escaped += char
                                else:
                                    escaped += f"\\u{cp}?"
                                    
                            prefix = ""
                            suffix = ""
                            if is_bold:
                                prefix += "\\b "
                                suffix += "\\b0 "
                            if is_italic:
                                prefix += "\\i "
                                suffix += "\\i0 "
                                
                            line_text += f"{prefix}{escaped}{suffix}"
                        if line_text:
                            block_text += line_text + " "
                    if block_text:
                        rtf_content += f"\\par {block_text}\n"
        rtf_content += "}"
        doc.close()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(rtf_content)
            
        return {"success": True, "output_path": output_path, "filename": output_filename}
        
    except Exception as e:
        print(f"[PDF→RTF] PyMuPDF failed: {e}. Falling back to plain text RTF.")
        
        text_content = _extract_pdf_text_with_ocr_fallback(input_path)
        if not text_content.strip():
            return {"success": False, "error": "Could not extract text from PDF."}

        rtf_content = "{\\rtf1\\ansi\\deff0\n"
        for para in text_content.split('\n'):
            if para.strip():
                escaped = ""
                for char in para:
                    cp = ord(char)
                    if cp < 128:
                        if char in ['\\', '{', '}']:
                            escaped += '\\' + char
                        else:
                            escaped += char
                    else:
                        escaped += f"\\u{cp}?"
                rtf_content += f"\\par {escaped}\n"
        rtf_content += "}"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(rtf_content)
        
        return {"success": True, "output_path": output_path, "filename": output_filename}


def _pdf_to_images(input_path: str, output_dir: str, target_format: str, name: str) -> dict:
    """Convert PDF pages to images using PyMuPDF (no Poppler needed)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {"success": False, "error": "PyMuPDF not installed. Run 'pip install PyMuPDF'"}
    
    try:
        pdf_doc = fitz.open(input_path)
        output_files = []
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            
            img_filename = f"{name}_page_{page_num + 1}.{target_format}"
            img_path = os.path.join(output_dir, img_filename)
            
            if target_format.lower() in ['jpg', 'jpeg']:
                pix.save(img_path, output="jpeg", jpg_quality=95)
            else:
                pix.save(img_path)
            
            output_files.append(img_filename)
        
        pdf_doc.close()
        
        if output_files:
            return {
                "success": True, 
                "output_path": os.path.join(output_dir, output_files[0]), 
                "filename": output_files[0],
                "all_files": output_files,
                "note": f"Created {len(output_files)} page images"
            }
        
        return {"success": False, "error": "Could not convert page"}
        
    except Exception as e:
        return {"success": False, "error": f"PDF to image conversion error: {str(e)}"}
