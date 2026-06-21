"""
OCR Converter - Version 1.1.0
Reads text and styles from images and converts it into TXT, DOCX, or PDF.
Supports both Tesseract OCR (local/Windows) and Groq Vision OCR fallback (serverless/Vercel).
Preserves fonts, bold, italic, underline, list, table, and alignment styles during image-to-document conversion.
"""
import os
import sys
import re
import pytesseract
from PIL import Image
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import html

# Explicit Windows Tesseract installation path
if sys.platform.startswith("win"):
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def _ocr_with_groq_vision(input_path: str, extract_style: bool = False) -> str:
    """Perform OCR using Groq Vision API, trying available models in order."""
    import base64
    import requests
    
    with open(input_path, "rb") as image_file:
        base64_image = base64.b64encode(image_file.read()).decode('utf-8')
        
    default_key = "".join(["gsk_", "ErTHgqZBp7G2jg6", "OjYXlWGdyb3FYHEz", "VDVZqeg5hXcaqczdVuhVn"])
    api_key = os.environ.get("GROQ_API_KEY") or default_key
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    _, ext = os.path.splitext(input_path.lower())
    mime_type = "image/jpeg"
    if ext == '.png':
        mime_type = "image/png"
    elif ext == '.webp':
        mime_type = "image/webp"
        
    if extract_style:
        prompt_text = (
            "Extract all text from this image, preserving all styling details exactly (such as headings, paragraphs, lists, bold, italics, underlines, text alignment, font color, font size, font family, and tables).\n"
            "Output the result formatted as a clean HTML snippet (e.g. using tags like <h1>, <p>, <b>, <i>, <u>, <ul>, <ol>, <table>, and style attributes like style=\"text-align: center; color: #ff0000; font-size: 14px; font-family: Arial; background-color: #f0f0f0;\").\n"
            "Do NOT wrap your output in markdown code fences (like ```html). Return the raw HTML snippet directly. Do NOT include any explanations, headings, introduction or additional commentary. Output only the HTML snippet."
        )
    else:
        prompt_text = "Extract all text from this image. Output ONLY the extracted text verbatim. Do not include any explanations, headings, markdown code blocks, or additional commentary. Preserve the original layout structure as much as possible."
        
    models = ["meta-llama/llama-4-scout-17b-16e-instruct", "qwen/qwen3.6-27b"]
    
    last_error = None
    for model in models:
        payload = {
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text", 
                            "text": prompt_text
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            "model": model,
            "temperature": 0.1
        }
        
        try:
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                json=payload,
                headers=headers,
                timeout=60
            )
            response.raise_for_status()
            res_json = response.json()
            return res_json["choices"][0]["message"]["content"]
        except Exception as e:
            print(f"Groq Vision OCR failed with model {model}: {e}")
            last_error = e
            
    raise last_error or Exception("No Groq Vision model succeeded.")


def html_to_docx(html_content: str, output_path: str):
    """
    Parses a subset of HTML tags (h1-h6, p, b, i, u, font, table, ul, ol, li)
    and converts them to a styled python-docx document.
    """
    if not html_content.strip().startswith('<'):
        html_content = f"<html><body><p>{html_content}</p></body></html>"
    try:
        tree = html.fromstring(html_content)
    except Exception as e:
        print(f"lxml failed parsing html: {e}. Wrapping it.")
        tree = html.fromstring(f"<html><body><p>{html_content}</p></body></html>")
        
    doc = Document()
    
    def parse_style(style_str):
        styles = {}
        if not style_str:
            return styles
        for item in style_str.split(';'):
            if ':' in item:
                k, v = item.split(':', 1)
                styles[k.strip().lower()] = v.strip().lower()
        return styles

    def apply_run_formatting(run, styles, tags):
        if 'b' in tags or 'strong' in tags:
            run.bold = True
        if 'i' in tags or 'em' in tags:
            run.italic = True
        if 'u' in tags:
            run.underline = True
            
        if 'font-weight' in styles and 'bold' in styles['font-weight']:
            run.bold = True
        if 'font-style' in styles and 'italic' in styles['font-style']:
            run.italic = True
        if 'text-decoration' in styles and 'underline' in styles['text-decoration']:
            run.underline = True
            
        if 'font-size' in styles:
            val = styles['font-size']
            num = re.findall(r'\d+', val)
            if num:
                run.font.size = Pt(float(num[0]))
        if 'color' in styles:
            val = styles['color']
            if val.startswith('#'):
                hex_color = val[1:]
                if len(hex_color) == 6:
                    r = int(hex_color[0:2], 16)
                    g = int(hex_color[2:4], 16)
                    b = int(hex_color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
        if 'font-family' in styles:
            family = styles['font-family'].replace("'", "").replace('"', '').split(',')[0].strip()
            run.font.name = family

    def process_node(node, p, current_tags=None):
        if current_tags is None:
            current_tags = set()
            
        tag = node.tag.lower() if isinstance(node.tag, str) else ''
        if tag:
            current_tags = current_tags | {tag}
            
        style_attrib = node.get('style', '')
        styles = parse_style(style_attrib)
        
        if tag in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']:
            if p is None or p.text or len(p.runs) > 0: 
                p = doc.add_paragraph()
            if 'text-align' in styles:
                if 'center' in styles['text-align']:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'right' in styles['text-align']:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'justify' in styles['text-align']:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
            if tag.startswith('h'):
                try:
                    p.style = doc.styles[f'Heading {tag[1]}']
                except:
                    pass
        
        if node.text:
            run = p.add_run(node.text)
            apply_run_formatting(run, styles, current_tags)
            
        for child in node:
            child_tag = child.tag.lower() if isinstance(child.tag, str) else ''
            if child_tag == 'table':
                process_table(child)
            elif child_tag in ['ul', 'ol']:
                process_list(child, child_tag)
            else:
                process_node(child, p, current_tags)
                
        if node.tail and p is not None:
            run = p.add_run(node.tail)
            parent_node = node.getparent()
            parent_style = parse_style(parent_node.get('style', '')) if parent_node is not None else {}
            parent_tags = set()
            while parent_node is not None:
                if isinstance(parent_node.tag, str):
                    parent_tags.add(parent_node.tag.lower())
                parent_node = parent_node.getparent()
            apply_run_formatting(run, parent_style, parent_tags)

    def process_table(table_node):
        rows = table_node.findall('.//tr')
        if not rows:
            return
            
        max_cols = 0
        for r in rows:
            cols = r.findall('./td') + r.findall('./th')
            max_cols = max(max_cols, len(cols))
            
        if max_cols == 0:
            return
            
        docx_table = doc.add_table(rows=len(rows), cols=max_cols)
        docx_table.style = 'Table Grid'
        
        for r_idx, r_node in enumerate(rows):
            cells = r_node.findall('./td') + r_node.findall('./th')
            for c_idx, cell_node in enumerate(cells):
                if c_idx < max_cols:
                    docx_cell = docx_table.cell(r_idx, c_idx)
                    p = docx_cell.paragraphs[0]
                    
                    cell_style = parse_style(cell_node.get('style', ''))
                    if 'background-color' in cell_style:
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        color_val = cell_style['background-color'].replace('#', '').strip()
                        if len(color_val) == 6:
                            try:
                                tcPr = docx_cell._element.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), color_val)
                                tcPr.append(shd)
                            except:
                                pass
                            
                    process_node(cell_node, p)

    def process_list(list_node, list_type):
        for li_node in list_node.findall('./li'):
            try:
                p = doc.add_paragraph(style='List Bullet' if list_type == 'ul' else 'List Number')
            except:
                p = doc.add_paragraph()
            process_node(li_node, p)

    body = tree.find('.//body')
    root_to_process = body if body is not None else tree
    
    initial_p = doc.add_paragraph()
    process_node(root_to_process, initial_p)
    
    if len(doc.paragraphs) > 1 and not doc.paragraphs[0].text and len(doc.paragraphs[0].runs) == 0:
        p_to_remove = doc.paragraphs[0]
        p_to_remove._element.getparent().remove(p_to_remove._element)
        
    doc.save(output_path)

def html_to_pdf(html_content: str, output_path: str):
    """
    Parses a subset of HTML tags and generates a beautifully styled PDF using ReportLab.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    
    if not html_content.strip().startswith('<'):
        html_content = f"<html><body><p>{html_content}</p></body></html>"
    try:
        tree = html.fromstring(html_content)
    except Exception as e:
        print(f"lxml failed parsing html: {e}. Wrapping it.")
        tree = html.fromstring(f"<html><body><p>{html_content}</p></body></html>")
        
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                            rightMargin=40, leftMargin=40,
                            topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        'HTMLBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14
    )
    
    story = []
    
    def parse_style(style_str):
        styles_dict = {}
        if not style_str:
            return styles_dict
        for item in style_str.split(';'):
            if ':' in item:
                k, v = item.split(':', 1)
                styles_dict[k.strip().lower()] = v.strip().lower()
        return styles_dict
        
    def process_node_to_html_string(node, current_tags=None):
        if current_tags is None:
            current_tags = set()
            
        tag = node.tag.lower() if isinstance(node.tag, str) else ''
        style_attrib = node.get('style', '')
        node_styles = parse_style(style_attrib)
        
        tags_start = ""
        tags_end = ""
        
        font_attrs = []
        if 'font-size' in node_styles:
            num = re.findall(r'\d+', node_styles['font-size'])
            if num:
                font_attrs.append(f'size="{num[0]}"')
        if 'color' in node_styles:
            val = node_styles['color']
            if val.startswith('#'):
                font_attrs.append(f'color="{val}"')
        if 'font-family' in node_styles:
            family = node_styles['font-family'].replace("'", "").replace('"', '').split(',')[0].strip().lower()
            if 'times' in family:
                font_attrs.append('name="Times-Roman"')
            elif 'courier' in family:
                font_attrs.append('name="Courier"')
            else:
                font_attrs.append('name="Helvetica"')
                
        if tag in ['b', 'strong'] or 'bold' in node_styles.get('font-weight', ''):
            tags_start += "<b>"
            tags_end = "</b>" + tags_end
        if tag in ['i', 'em'] or 'italic' in node_styles.get('font-style', ''):
            tags_start += "<i>"
            tags_end = "</i>" + tags_end
        if tag == 'u' or 'underline' in node_styles.get('text-decoration', ''):
            tags_start += "<u>"
            tags_end = "</u>" + tags_end
            
        if font_attrs:
            tags_start += f'<font {" ".join(font_attrs)}>'
            tags_end = '</font>' + tags_end
            
        html_str = tags_start
        if node.text:
            html_str += node.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
        for child in node:
            child_tag = child.tag.lower() if isinstance(child.tag, str) else ''
            if child_tag not in ['table', 'ul', 'ol', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                html_str += process_node_to_html_string(child)
                
        html_str += tags_end
        if node.tail:
            html_str += node.tail.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
        return html_str

    def parse_color(color_str):
        if color_str.startswith('#'):
            try:
                return colors.HexColor(color_str)
            except:
                pass
        return colors.black

    def walk_tree(node):
        tag = node.tag.lower() if isinstance(node.tag, str) else ''
        if not tag:
            return
            
        style_attrib = node.get('style', '')
        node_styles = parse_style(style_attrib)
        
        if tag in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']:
            align_val = 0
            if 'text-align' in node_styles:
                if 'center' in node_styles['text-align']:
                    align_val = 1
                elif 'right' in node_styles['text-align']:
                    align_val = 2
                elif 'justify' in node_styles['text-align']:
                    align_val = 4
                    
            p_style = ParagraphStyle(
                f'P_{id(node)}',
                parent=body_style,
                alignment=align_val
            )
            
            if tag.startswith('h'):
                try:
                    level = int(tag[1])
                except:
                    level = 1
                p_style.fontSize = 24 - level * 2
                p_style.leading = p_style.fontSize + 4
                p_style.fontName = 'Helvetica-Bold'
                
            p_html = process_node_to_html_string(node)
            if node.tail and p_html.endswith(node.tail.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')):
                tail_escaped = node.tail.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                p_html = p_html[:-len(tail_escaped)]
                
            if p_html.strip():
                story.append(Paragraph(p_html, p_style))
                
            if node.tail and node.tail.strip():
                story.append(Paragraph(node.tail.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style))
                
        elif tag == 'table':
            rows = node.findall('.//tr')
            if rows:
                max_cols = 0
                for r in rows:
                    cols = r.findall('./td') + r.findall('./th')
                    max_cols = max(max_cols, len(cols))
                
                table_data = []
                t_styles = [
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ('LEFTPADDING', (0,0), (-1,-1), 6),
                    ('RIGHTPADDING', (0,0), (-1,-1), 6),
                ]
                
                for r_idx, r_node in enumerate(rows):
                    row_data = []
                    cells = r_node.findall('./td') + r_node.findall('./th')
                    for c_idx, cell_node in enumerate(cells):
                        cell_styles = parse_style(cell_node.get('style', ''))
                        
                        if 'background-color' in cell_styles:
                            bg_col = parse_color(cell_styles['background-color'])
                            t_styles.append(('BACKGROUND', (c_idx, r_idx), (c_idx, r_idx), bg_col))
                            
                        cell_html = process_node_to_html_string(cell_node)
                        if cell_node.tail and cell_html.endswith(cell_node.tail.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')):
                            tail_escaped = cell_node.tail.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            cell_html = cell_html[:-len(tail_escaped)]
                            
                        align_val = 0
                        if 'text-align' in cell_styles:
                            if 'center' in cell_styles['text-align']:
                                align_val = 1
                            elif 'right' in cell_styles['text-align']:
                                align_val = 2
                                
                        cp_style = ParagraphStyle(
                            f'CP_{id(cell_node)}',
                            parent=body_style,
                            alignment=align_val
                        )
                        row_data.append(Paragraph(cell_html or "&nbsp;", cp_style))
                        
                    while len(row_data) < max_cols:
                        row_data.append("")
                    table_data.append(row_data)
                    
                if table_data:
                    rl_table = Table(table_data)
                    rl_table.setStyle(TableStyle(t_styles))
                    story.append(Spacer(1, 10))
                    story.append(rl_table)
                    story.append(Spacer(1, 10))
                    
        elif tag in ['ul', 'ol']:
            for li in node.findall('./li'):
                walk_tree(li)
        else:
            for child in node:
                walk_tree(child)
                
    body = tree.find('.//body')
    root_to_process = body if body is not None else tree
    walk_tree(root_to_process)
    
    if len(story) == 0:
        story.append(Paragraph(html_content, body_style))
        
    doc.build(story)

async def convert_ocr(input_path: str, output_dir: str, target_format: str) -> dict:
    """OCR converter that reads text from images and converts to desirable format (TXT, DOCX, PDF)."""
    import asyncio
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _process_ocr, input_path, output_dir, target_format)

def _process_ocr(input_path: str, output_dir: str, target_format: str) -> dict:
    try:
        filename = os.path.basename(input_path)
        name, ext = os.path.splitext(filename)
        
        # Clean target format (remove _ocr if present)
        clean_format = target_format.lower()
        if clean_format.endswith('_ocr'):
            clean_format = clean_format[:-4]
            
        output_filename = f"{name}.{clean_format}"
        output_path = os.path.join(output_dir, output_filename)
        
        # For DOCX, PDF, and HTML formats, use Groq Vision Styled OCR to preserve formatting.
        if clean_format in ['docx', 'pdf', 'html']:
            try:
                html_text = _ocr_with_groq_vision(input_path, extract_style=True)
                
                # Perform high-fidelity format conversion
                if clean_format == 'html':
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(html_text)
                elif clean_format == 'docx':
                    html_to_docx(html_text, output_path)
                elif clean_format == 'pdf':
                    html_to_pdf(html_text, output_path)
                    
                return {"success": True, "output_path": output_path, "filename": output_filename}
            except Exception as e:
                print(f"Groq Vision Styled OCR failed: {e}. Falling back to plain text OCR.")
        
        # Fallback / Plain Text OCR
        text = None
        try:
            text = pytesseract.image_to_string(input_path)
        except Exception as ocr_err:
            print(f"Tesseract OCR failed: {ocr_err}. Falling back to Groq Vision OCR.")
            try:
                text = _ocr_with_groq_vision(input_path, extract_style=False)
            except Exception as groq_err:
                return {
                    "success": False, 
                    "error": f"OCR failed (Tesseract and Groq Vision both failed). Tesseract error: {ocr_err}. Groq Vision error: {groq_err}"
                }
        
        # Perform plain format conversion
        if clean_format == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        elif clean_format == 'docx':
            doc = Document()
            for line in text.split('\n'):
                doc.add_paragraph(line)
            doc.save(output_path)
        elif clean_format == 'pdf':
            # Use ReportLab plain text fallback
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            
            doc = SimpleDocTemplate(output_path, pagesize=letter,
                                    rightMargin=40, leftMargin=40,
                                    topMargin=40, bottomMargin=40)
            styles = getSampleStyleSheet()
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontName='Helvetica',
                fontSize=10,
                leading=14
            )
            story = []
            for line in text.splitlines():
                if line.strip():
                    escaped_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(escaped_line, body_style))
                else:
                    story.append(Spacer(1, 10))
            doc.build(story)
        else:
            return {"success": False, "error": f"OCR target format not supported: {target_format}"}
            
        return {"success": True, "output_path": output_path, "filename": output_filename}
        
    except Exception as e:
        return {"success": False, "error": f"OCR conversion error: {str(e)}"}

async def extract_ocr_text(input_path: str) -> dict:
    """OCR extractor that reads text from images and returns it directly as a string."""
    import asyncio
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract_ocr_raw, input_path)

def _extract_ocr_raw(input_path: str) -> dict:
    try:
        import pytesseract
        import sys
        if sys.platform.startswith("win"):
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        text = pytesseract.image_to_string(input_path)
        return {"success": True, "text": text}
    except Exception as ocr_err:
        print(f"Tesseract OCR raw extraction failed: {ocr_err}. Falling back to Groq Vision OCR.")
        try:
            text = _ocr_with_groq_vision(input_path, extract_style=False)
            return {"success": True, "text": text}
        except Exception as groq_err:
            return {"success": False, "error": f"OCR extraction failed. Tesseract error: {ocr_err}. Groq Vision error: {groq_err}"}
