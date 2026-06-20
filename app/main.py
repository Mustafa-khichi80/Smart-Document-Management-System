"""
Smart Document Management - Main API
FastAPI backend for document operations and AI editing.
"""
from fastapi import FastAPI, File, UploadFile
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response
import shutil
import os
import time
import threading
from .utils import (
    get_output_dir,
    clean_filename,
    apply_pdf_watermark,
    apply_docx_watermark,
    apply_pptx_watermark,
    apply_image_watermark,
    apply_html_watermark
)
from .converters import (
    convert_image,
    convert_doc,
    convert_pdf,
    convert_docx,
    convert_pptx,
    convert_archive,
    convert_ocr,
    extract_ocr_text
)
from .converters.ai import edit_document_with_ai
from pydantic import BaseModel
import webbrowser


def get_file_extension(filename: str) -> str:
    """Get file extension, handling double extensions like .tar.gz"""
    filename_lower = filename.lower()
    # Check for double extensions first
    double_exts = ['.tar.gz', '.tar.bz2', '.tar.xz']
    for ext in double_exts:
        if filename_lower.endswith(ext):
            return ext
    # Otherwise use standard splitext
    return os.path.splitext(filename)[1].lower()

class ConvertRequest(BaseModel):
    file_path: str
    target_format: str
    quality: str = "high"

class AIEditRequest(BaseModel):
    text: str
    prompt: str
    api_key: str = None
    filename: str = None

class OCRExtractRequest(BaseModel):
    file_path: str

class SaveDocumentRequest(BaseModel):
    text: str
    filename: str
    target_format: str

app = FastAPI(title="Smart Document Management")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
if not os.path.exists(static_dir):
    os.makedirs(static_dir)
app.mount("/static", StaticFiles(directory=static_dir), name="static")

@app.get("/")
async def read_index():
    return FileResponse(
        os.path.join(static_dir, 'index.html'),
        headers={"Cache-Control": "no-store, no-cache, must-revalidate, max-age=0"}
    )

@app.get('/favicon.ico', include_in_schema=False)
async def favicon():
    return Response(status_code=204)


@app.get("/api/languages")
async def get_languages():
    """Scan locales directory and return available languages."""
    locales_dir = os.path.join(static_dir, "locales")
    languages = []
    if os.path.exists(locales_dir):
        for f in os.listdir(locales_dir):
            if f.endswith(".json"):
                lang_code = os.path.splitext(f)[0]
                languages.append(lang_code)
    languages.sort()
    return {"languages": languages}

IS_VERCEL = os.environ.get("VERCEL") == "1"
if IS_VERCEL:
    UPLOAD_DIR = "/tmp/temp_uploads"
else:
    UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "temp_uploads")

if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """Save uploaded file temporarily and analyze its type."""
    try:
        safe_filename = clean_filename(file.filename)
        file_path = os.path.join(UPLOAD_DIR, safe_filename)
        
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        ext = get_file_extension(safe_filename)
        size = os.path.getsize(file_path)
        
        file_type = "unknown"
        if ext in ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.tif', '.ico', '.gif', '.heic', '.heif', '.svg', '.avif']:
            file_type = "image"

        elif ext in ['.csv', '.xlsx', '.xls', '.json', '.xml', '.html', '.txt']:
            file_type = "data"
        elif ext == '.pdf':
            file_type = "pdf"
        elif ext in ['.docx', '.doc']:
            file_type = "docx"
        elif ext in ['.pptx', '.ppt']:
            file_type = "pptx"
        elif ext in ['.zip', '.7z', '.tar', '.gz', '.tgz', '.bz2', '.tar.gz', '.tar.bz2', '.tar.xz']:
            file_type = "archive"

        return {
            "filename": safe_filename,
            "original_name": file.filename,
            "path": file_path,
            "type": file_type,
            "size": size,
            "extension": ext
        }
    except Exception as e:
         return JSONResponse(status_code=500, content={"message": str(e)})

@app.on_event("shutdown")
def remove_temp_files():
    if os.path.exists(UPLOAD_DIR):
        try:
            shutil.rmtree(UPLOAD_DIR)
            os.makedirs(UPLOAD_DIR)
        except:
            pass

@app.on_event("startup")
async def startup_event():
    """Open browser on startup and start cleanup thread."""
    if os.environ.get("VERCEL") == "1":
        return
        
    def open_browser():
        time.sleep(1.5)
        webbrowser.open("http://localhost:1453")
        
    threading.Thread(target=open_browser, daemon=True).start()
    
    def cleanup_old_files():
        while True:
            time.sleep(600)
            try:
                now = time.time()
                if os.path.exists(UPLOAD_DIR):
                    for f in os.listdir(UPLOAD_DIR):
                        fpath = os.path.join(UPLOAD_DIR, f)
                        if os.path.isfile(fpath):
                            age = now - os.path.getmtime(fpath)
                            if age > 600:
                                os.remove(fpath)
                                print(f"[Cleanup] Deleted: {f}")
            except Exception as e:
                print(f"[Cleanup] Error: {e}")
    
    cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
    cleanup_thread.start()

@app.post("/api/convert")
async def api_convert(request: ConvertRequest):
    output_dir = get_output_dir()
    file_path = os.path.join(UPLOAD_DIR, request.file_path)
    
    if not os.path.exists(file_path):
        return {"success": False, "error": f"File not found: {request.file_path}"}
    
    ext = get_file_extension(os.path.basename(file_path))
    result = {"success": False, "error": "Unknown file type"}

    try:
        # Image formats
        if ext in ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.tif', '.ico', '.gif', '.heic', '.heif', '.svg', '.avif']:
            if request.target_format in ['txt', 'docx', 'txt_ocr', 'docx_ocr', 'pdf_ocr']:
                result = await convert_ocr(file_path, output_dir, request.target_format)
            else:
                result = await convert_image(file_path, output_dir, request.target_format, request.quality)

        # Data formats
        elif ext in ['.csv', '.xlsx', '.xls', '.json', '.xml', '.html', '.txt']:
            result = await convert_doc(file_path, output_dir, request.target_format)
        # PDF
        elif ext == '.pdf':
            result = await convert_pdf(file_path, output_dir, request.target_format)
        # Word documents
        elif ext in ['.docx', '.doc']:
            result = await convert_docx(file_path, output_dir, request.target_format)
        # PowerPoint
        elif ext in ['.pptx', '.ppt']:
            result = await convert_pptx(file_path, output_dir, request.target_format)
        # Archives
        elif ext in ['.zip', '.7z', '.tar', '.gz', '.tgz', '.bz2', '.tar.gz', '.tar.bz2', '.tar.xz']:
            result = await convert_archive(file_path, output_dir, request.target_format)
            
        # Apply watermarks on success
        if result.get("success"):
            output_path = result.get("output_path")
            all_files = result.get("all_files", [])
            
            files_to_watermark = []
            if all_files:
                for f in all_files:
                    files_to_watermark.append(os.path.join(output_dir, f))
            elif output_path:
                files_to_watermark.append(output_path)
                
            for path in files_to_watermark:
                if os.path.exists(path):
                    _, out_ext = os.path.splitext(path.lower())
                    if out_ext == '.pdf':
                        apply_pdf_watermark(path)
                    elif out_ext == '.docx':
                        apply_docx_watermark(path)
                    elif out_ext == '.pptx':
                        apply_pptx_watermark(path)
                    elif out_ext in ['.png', '.jpg', '.jpeg', '.webp', '.bmp', '.tiff', '.tif', '.gif']:
                        apply_image_watermark(path)
                    elif out_ext in ['.html', '.htm']:
                        apply_html_watermark(path)
                        
    except Exception as e:
        result = {"success": False, "error": f"Dönüşüm hatası: {str(e)}"}
    
    return result

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    """Download converted file."""
    safe_filename = os.path.basename(filename)
    file_path = os.path.join(get_output_dir(), safe_filename)
    
    if os.path.exists(file_path):
        return FileResponse(
            path=file_path,
            filename=safe_filename,
            media_type='application/octet-stream'
        )
    return JSONResponse(status_code=404, content={"error": "File not found"})

@app.post("/api/download-all")
async def download_all_files(request: dict):
    """Download multiple files as ZIP."""
    import zipfile
    import io
    
    filenames = request.get('filenames', [])
    if not filenames:
        return JSONResponse(status_code=400, content={"error": "File list empty"})

    output_dir = get_output_dir()
    zip_buffer = io.BytesIO()

    try:
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for filename in filenames:
                safe_name = os.path.basename(filename)
                file_path = os.path.join(output_dir, safe_name)
                if os.path.exists(file_path):
                    zip_file.write(file_path, safe_name)
        
        zip_buffer.seek(0)
        return Response(
            content=zip_buffer.getvalue(),
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=downloads.zip"}
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/api/ocr/extract")
async def api_ocr_extract(request: OCRExtractRequest):
    file_path = os.path.join(UPLOAD_DIR, os.path.basename(request.file_path))
    if not os.path.exists(file_path):
        return {"success": False, "error": f"Image file not found: {request.file_path}"}
    
    result = await extract_ocr_text(file_path)
    return result

@app.post("/api/ai/edit")
async def api_ai_edit(request: AIEditRequest):
    result = await edit_document_with_ai(request.text, request.prompt, request.api_key)
    return result

@app.post("/api/document/text")
async def api_document_text(request: OCRExtractRequest):
    file_path = os.path.join(UPLOAD_DIR, os.path.basename(request.file_path))
    if not os.path.exists(file_path):
        return {"success": False, "error": f"File not found: {request.file_path}"}
        
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    try:
        if ext in ['.txt', '.xml', '.html', '.csv', '.json']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
                
        elif ext in ['.docx', '.doc']:
            from docx import Document
            doc = Document(file_path)
            text_lines = []
            for p in doc.paragraphs:
                text_lines.append(p.text)
            for t in doc.tables:
                for r in t.rows:
                    text_lines.append("\t".join([cell.text for cell in r.cells]))
            text = "\n".join(text_lines)
            
        elif ext == '.pdf':
            from .converters.pdf import _extract_pdf_text_with_ocr_fallback
            text = _extract_pdf_text_with_ocr_fallback(file_path)
            
        elif ext in ['.png', '.jpg', '.jpeg', '.webp', '.bmp', '.tiff', '.tif']:
            res = await extract_ocr_text(file_path)
            if res.get("success"):
                text = res.get("text")
            else:
                return {"success": False, "error": res.get("error")}
                
        else:
            return {"success": False, "error": f"Cannot load text for file extension: {ext}"}
            
        return {"success": True, "text": text}
        
    except Exception as e:
        return {"success": False, "error": f"Failed to read document: {str(e)}"}

def update_paragraph_text_preserving_runs(p, new_text):
    """
    Updates the text of a paragraph while preserving formatting (bold, italic, underline, 
    font family, size, color) at character/run boundaries using character diff alignment.
    """
    import difflib
    
    old_text = "".join(r.text for r in p.runs)
    if old_text == new_text:
        return
        
    def copy_run_format(src, dest):
        try:
            dest.bold = src.bold
        except:
            pass
        try:
            dest.italic = src.italic
        except:
            pass
        try:
            dest.underline = src.underline
        except:
            pass
        try:
            if src.font.name:
                dest.font.name = src.font.name
        except:
            pass
        try:
            if src.font.size:
                dest.font.size = src.font.size
        except:
            pass
        try:
            if src.font.color and src.font.color.rgb:
                dest.font.color.rgb = src.font.color.rgb
        except:
            pass

    # Build character-to-run mapping for old text
    old_char_runs = []
    for r in p.runs:
        for _ in r.text:
            old_char_runs.append(r)
            
    matcher = difflib.SequenceMatcher(None, old_text, new_text)
    new_char_runs = []
    
    for op, s1, e1, s2, e2 in matcher.get_opcodes():
        if op == 'equal':
            for idx in range(s1, e1):
                new_char_runs.append((old_text[idx], old_char_runs[idx]))
        elif op in ('replace', 'insert'):
            # Determine source run style for the new/modified characters
            if s1 > 0:
                source_run = old_char_runs[s1 - 1]
            elif s1 < len(old_char_runs):
                source_run = old_char_runs[s1]
            elif p.runs:
                source_run = p.runs[0]
            else:
                source_run = None
                
            for idx in range(s2, e2):
                new_char_runs.append((new_text[idx], source_run))

    # Group characters by source run
    grouped_runs = []
    if new_char_runs:
        current_run_source = new_char_runs[0][1]
        current_text = [new_char_runs[0][0]]
        
        for char, run in new_char_runs[1:]:
            if run is current_run_source:
                current_text.append(char)
            else:
                grouped_runs.append(("".join(current_text), current_run_source))
                current_run_source = run
                current_text = [char]
        grouped_runs.append(("".join(current_text), current_run_source))

    # Clear old runs from paragraph element
    p_element = p._element
    for r in list(p.runs):
        p_element.remove(r._r)
        
    # Recreate runs with their respective formatting
    for text_segment, source_run in grouped_runs:
        new_run = p.add_run(text_segment)
        if source_run is not None:
            copy_run_format(source_run, new_run)

@app.post("/api/document/save")
async def api_document_save(request: SaveDocumentRequest):
    output_dir = get_output_dir()
    name, _ = os.path.splitext(os.path.basename(request.filename))
    target_format = request.target_format.lower().strip('.')
    
    output_filename = f"{name}_edited.{target_format}"
    output_path = os.path.join(output_dir, output_filename)
    
    try:
        if target_format == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(request.text)
                
        elif target_format == 'xml':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(request.text)
                
        elif target_format == 'docx':
            original_path = os.path.join(UPLOAD_DIR, os.path.basename(request.filename))
            if os.path.exists(original_path) and original_path.lower().endswith('.docx'):
                from docx import Document
                doc = Document(original_path)
                
                lines = request.text.split('\n')
                line_idx = 0
                for p in doc.paragraphs:
                    if line_idx < len(lines):
                        update_paragraph_text_preserving_runs(p, lines[line_idx])
                        line_idx += 1
                for t in doc.tables:
                    for r in t.rows:
                        if line_idx < len(lines):
                            row_text = lines[line_idx]
                            cell_texts = row_text.split('\t')
                            for c_idx, cell in enumerate(r.cells):
                                if c_idx < len(cell_texts):
                                    if cell.paragraphs:
                                        update_paragraph_text_preserving_runs(cell.paragraphs[0], cell_texts[c_idx])
                                        for extra_p in list(cell.paragraphs[1:]):
                                            extra_p._element.getparent().remove(extra_p._element)
                                    else:
                                        cell.text = cell_texts[c_idx]
                            line_idx += 1
                while line_idx < len(lines):
                    doc.add_paragraph(lines[line_idx])
                    line_idx += 1
                    
                doc.save(output_path)
            else:
                from docx import Document
                doc = Document()
                for line in request.text.split('\n'):
                    doc.add_paragraph(line)
                doc.save(output_path)
            
        elif target_format == 'pdf':
            original_path = os.path.join(UPLOAD_DIR, os.path.basename(request.filename))
            if os.path.exists(original_path) and original_path.lower().endswith('.docx'):
                from docx import Document
                doc = Document(original_path)
                
                lines = request.text.split('\n')
                line_idx = 0
                for p in doc.paragraphs:
                    if line_idx < len(lines):
                        update_paragraph_text_preserving_runs(p, lines[line_idx])
                        line_idx += 1
                for t in doc.tables:
                    for r in t.rows:
                        if line_idx < len(lines):
                            row_text = lines[line_idx]
                            cell_texts = row_text.split('\t')
                            for c_idx, cell in enumerate(r.cells):
                                if c_idx < len(cell_texts):
                                    if cell.paragraphs:
                                        update_paragraph_text_preserving_runs(cell.paragraphs[0], cell_texts[c_idx])
                                        for extra_p in list(cell.paragraphs[1:]):
                                            extra_p._element.getparent().remove(extra_p._element)
                                    else:
                                        cell.text = cell_texts[c_idx]
                            line_idx += 1
                while line_idx < len(lines):
                    doc.add_paragraph(lines[line_idx])
                    line_idx += 1
                    
                temp_docx_path = output_path + ".docx"
                doc.save(temp_docx_path)
                
                from .converters.docx_converter import _docx_to_pdf
                pdf_res = _docx_to_pdf(temp_docx_path, output_path, output_filename)
                
                if os.path.exists(temp_docx_path):
                    os.remove(temp_docx_path)
                    
                if not pdf_res.get("success"):
                    return {"success": False, "error": f"PDF conversion failed: {pdf_res.get('error')}"}
            else:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                
                doc = SimpleDocTemplate(output_path, pagesize=letter,
                                        rightMargin=40, leftMargin=40,
                                        topMargin=40, bottomMargin=40)
                styles = getSampleStyleSheet()
                body_style = ParagraphStyle(
                    'CustomBody',
                    parent=styles['BodyText'],
                    fontName='Helvetica',
                    fontSize=10,
                    leading=14
                )
                
                story = []
                for line in request.text.splitlines():
                    if not line.strip():
                        story.append(Spacer(1, 10))
                    else:
                        escaped_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(escaped_line, body_style))
                doc.build(story)
            
        else:
            return {"success": False, "error": f"Unsupported export format: {target_format}"}
            
        # Apply watermarking if active on output
        if os.path.exists(output_path):
            if target_format == 'pdf':
                apply_pdf_watermark(output_path)
            elif target_format == 'docx':
                apply_docx_watermark(output_path)
                
        return {"success": True, "filename": output_filename, "url": f"/api/download/{output_filename}"}
        
    except Exception as e:
        return {"success": False, "error": f"Failed to save document: {str(e)}"}
